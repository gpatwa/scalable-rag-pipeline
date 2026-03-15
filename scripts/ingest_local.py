#!/usr/bin/env python3
"""
Lightweight local document ingestion (no Ray required).

Usage:
    python scripts/ingest_local.py <file_or_dir>      # ingest a file or directory
    python scripts/ingest_local.py --sample            # ingest built-in sample data

Supports: .txt, .md, .pdf, .docx, .png, .jpg, .jpeg, .gif, .webp
Performs:  parse → chunk → embed (Ollama/Gemini) → upsert Qdrant → extract entities → Neo4j

When MULTIMODAL_ENABLED=true:
  - Images are embedded via Gemini into the rag_multimodal collection
  - PDFs have images extracted and embedded alongside text
  - Standalone image files (.png, .jpg, etc.) are supported
"""
import asyncio
import argparse
import glob
import os
import sys
import uuid
import logging

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from services.api.app.config import settings

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
MULTIMODAL_ENABLED = os.getenv("MULTIMODAL_ENABLED", "false").lower() == "true"

# Extend supported extensions when multimodal is on
if MULTIMODAL_ENABLED:
    SUPPORTED_EXTENSIONS = SUPPORTED_EXTENSIONS | IMAGE_EXTENSIONS

# ── Sample data for quick testing ────────────────────────────────────────────
SAMPLE_DOCUMENTS = [
    {
        "filename": "company_overview.txt",
        "text": (
            "Acme Corp was founded in 2020 by Jane Smith and John Doe. "
            "The company is headquartered in San Francisco, California. "
            "Acme Corp specializes in building enterprise AI solutions for healthcare and finance. "
            "Their flagship product is MedAssist, an AI-powered clinical decision support system. "
            "MedAssist uses retrieval-augmented generation to provide evidence-based recommendations. "
            "The system processes over 10 million patient records daily. "
            "Acme Corp raised a $50 million Series B round led by Sequoia Capital in 2023. "
            "The company has over 200 employees across offices in San Francisco, New York, and London."
        ),
    },
    {
        "filename": "product_docs.txt",
        "text": (
            "MedAssist Technical Architecture: "
            "The system uses a microservices architecture deployed on AWS EKS. "
            "The retrieval pipeline combines vector search using Qdrant with knowledge graph queries via Neo4j. "
            "Embeddings are generated using the BGE-M3 model served through Ray Serve. "
            "The LLM backbone is Llama 3 70B, optimized with vLLM for low-latency inference. "
            "Data ingestion supports PDF, DOCX, and HTML formats. "
            "Documents are chunked using a sliding window approach with 512-token chunks and 50-token overlap. "
            "The system achieves sub-2-second end-to-end latency for queries. "
            "Authentication is handled via JWT tokens with role-based access control."
        ),
    },
    {
        "filename": "team_info.txt",
        "text": (
            "Acme Corp Leadership Team: "
            "Jane Smith is the CEO and co-founder. She previously worked at Google Brain. "
            "John Doe is the CTO and co-founder. He led the ML infrastructure team at Meta. "
            "Sarah Chen is the VP of Engineering. She joined from Amazon Web Services. "
            "The engineering team follows an agile methodology with two-week sprints. "
            "The company uses Python and TypeScript as primary programming languages. "
            "Key technologies include FastAPI, React, Kubernetes, and Terraform. "
            "The team conducts weekly architecture reviews and monthly tech talks."
        ),
    },
]


# ── Text Cleaning ────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Remove surrogate characters and other problematic Unicode from extracted text."""
    # Encode to utf-8 with surrogates replaced, then decode back
    cleaned = text.encode("utf-8", errors="replace").decode("utf-8")
    # Strip null bytes
    cleaned = cleaned.replace("\x00", "")
    return cleaned


# ── File Parsers ─────────────────────────────────────────────────────────────

def parse_pdf(filepath: str) -> str:
    """Extract text from a PDF file using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(filepath)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(text.strip())
    full_text = clean_text("\n\n".join(pages))
    logger.info(f"  Parsed PDF: {len(reader.pages)} pages, {len(full_text)} chars")
    return full_text


def parse_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = clean_text("\n\n".join(paragraphs))
    logger.info(f"  Parsed DOCX: {len(doc.paragraphs)} paragraphs, {len(full_text)} chars")
    return full_text


def parse_text(filepath: str) -> str:
    """Read plain text / markdown files."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    logger.info(f"  Parsed text: {len(text)} chars")
    return text


def parse_file(filepath: str) -> str:
    """Route to the correct parser based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return parse_pdf(filepath)
    elif ext == ".docx":
        return parse_docx(filepath)
    elif ext in (".txt", ".md"):
        return parse_text(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")


# ── Multimodal helpers ──────────────────────────────────────────────────────

def _get_gemini_client():
    """Lazy-init a google-genai Client for multimodal embedding."""
    if not hasattr(_get_gemini_client, "_client"):
        from google import genai
        api_key = settings.GOOGLE_API_KEY
        if not api_key:
            raise RuntimeError("GOOGLE_API_KEY required for multimodal ingestion")
        _get_gemini_client._client = genai.Client(api_key=api_key)
        _get_gemini_client._model = getattr(settings, "GEMINI_EMBED_MODEL", "gemini-embedding-2-preview")
        _get_gemini_client._dims = getattr(settings, "GEMINI_EMBED_DIMENSIONS", 768)
        logger.info(f"  Gemini client initialized (model={_get_gemini_client._model}, dims={_get_gemini_client._dims})")
    return _get_gemini_client._client


async def embed_image_gemini(image_bytes: bytes, mime_type: str) -> list[float]:
    """Embed a single image via Gemini."""
    from google.genai import types
    client = _get_gemini_client()
    part = types.Part.from_bytes(data=image_bytes, mime_type=mime_type)
    result = client.models.embed_content(
        model=_get_gemini_client._model,
        contents=part,
        config=types.EmbedContentConfig(
            task_type="RETRIEVAL_DOCUMENT",
            output_dimensionality=_get_gemini_client._dims,
        ),
    )
    return list(result.embeddings[0].values)


async def embed_texts_gemini(texts: list[str]) -> list[list[float]]:
    """Embed texts via Gemini in batches of 100 with rate-limit retry."""
    import time
    from google.genai import types
    client = _get_gemini_client()
    all_embeddings = []
    batch_size = 100
    total_batches = (len(texts) + batch_size - 1) // batch_size
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        batch_num = i // batch_size + 1
        for attempt in range(5):
            try:
                result = client.models.embed_content(
                    model=_get_gemini_client._model,
                    contents=batch,
                    config=types.EmbedContentConfig(
                        task_type="RETRIEVAL_DOCUMENT",
                        output_dimensionality=_get_gemini_client._dims,
                    ),
                )
                all_embeddings.extend([list(e.values) for e in result.embeddings])
                logger.info(f"    Batch {batch_num}/{total_batches}: embedded {len(batch)} chunks")
                break
            except Exception as e:
                if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                    wait = 45 * (attempt + 1)
                    logger.info(f"    Rate limited, waiting {wait}s before retry...")
                    time.sleep(wait)
                else:
                    raise
    return all_embeddings


def upsert_qdrant_multimodal(
    chunks: list[dict], embeddings: list[list[float]], filename: str, tenant_id: str = "default"
):
    """Insert multimodal chunks (text + image) into the rag_multimodal collection."""
    from qdrant_client import QdrantClient
    from qdrant_client.http import models

    collection = settings.MULTIMODAL_COLLECTION
    client = QdrantClient(host=settings.QDRANT_HOST, port=settings.QDRANT_PORT)

    points = []
    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        if emb is None:
            continue
        payload = {
            "text": chunk.get("text", ""),
            "content_type": chunk.get("content_type", "text"),
            "metadata": {"filename": filename, "chunk_index": i},
            "tenant_id": tenant_id,
        }
        if chunk.get("content_type") == "image":
            payload["image_mime_type"] = chunk.get("mime_type", "image/png")
        points.append(
            models.PointStruct(id=str(uuid.uuid4()), vector=emb, payload=payload)
        )

    batch_size = 100
    for i in range(0, len(points), batch_size):
        client.upsert(collection_name=collection, points=points[i : i + batch_size])
    logger.info(f"  Qdrant ({collection}): upserted {len(points)} vectors (tenant={tenant_id})")


# ── Chunking ─────────────────────────────────────────────────────────────────
def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> list[str]:
    """Simple character-based sliding window chunker."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


# ── Embedding ────────────────────────────────────────────────────────────────
async def _embed_one(client, text: str, idx: int, semaphore) -> tuple[int, list[float]]:
    """Embed a single text with concurrency control."""
    async with semaphore:
        resp = await client.post(
            settings.RAY_EMBED_ENDPOINT,
            json={"model": settings.EMBED_MODEL, "prompt": text},
        )
        resp.raise_for_status()
        return (idx, resp.json()["embedding"])


async def embed_texts(texts: list[str], concurrency: int = 10) -> list[list[float]]:
    """Embed texts via Ollama with concurrent requests for speed."""
    import httpx

    semaphore = asyncio.Semaphore(concurrency)
    embeddings = [None] * len(texts)
    done_count = 0

    async with httpx.AsyncClient(timeout=120.0) as client:
        # Process in batches to show progress and avoid memory spikes
        batch_size = 50
        for batch_start in range(0, len(texts), batch_size):
            batch_end = min(batch_start + batch_size, len(texts))
            batch_texts = texts[batch_start:batch_end]

            tasks = [
                _embed_one(client, text, batch_start + i, semaphore)
                for i, text in enumerate(batch_texts)
            ]
            results = await asyncio.gather(*tasks)

            for idx, emb in results:
                embeddings[idx] = emb
            done_count += len(batch_texts)
            logger.info(f"  Embedded {done_count}/{len(texts)} chunks...")

    return embeddings


# ── Qdrant upsert ────────────────────────────────────────────────────────────
def upsert_qdrant(chunks: list[str], embeddings: list[list[float]], filename: str, tenant_id: str = "default"):
    """Insert chunk vectors into Qdrant in batches, tagged with tenant_id."""
    from qdrant_client import QdrantClient
    from qdrant_client.http import models

    client = QdrantClient(host=settings.QDRANT_HOST, port=settings.QDRANT_PORT)

    points = [
        models.PointStruct(
            id=str(uuid.uuid4()),
            vector=emb,
            payload={
                "text": chunk,
                "metadata": {"filename": filename, "chunk_index": i},
                "tenant_id": tenant_id,
            },
        )
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings))
    ]

    # Batch upsert to avoid payload size issues with large documents
    batch_size = 100
    for i in range(0, len(points), batch_size):
        batch = points[i : i + batch_size]
        client.upsert(collection_name=settings.QDRANT_COLLECTION, points=batch)
    logger.info(f"  ✅ Qdrant: upserted {len(points)} vectors (tenant={tenant_id})")


# ── Neo4j entities ───────────────────────────────────────────────────────────
async def create_neo4j_entities(text: str, filename: str, tenant_id: str = "default"):
    """
    Extract simple entity triples using the LLM, then MERGE into Neo4j.
    Falls back to inserting the document as a single node if extraction fails.
    """
    import httpx
    from neo4j import GraphDatabase

    # Ask the LLM to extract triples
    extraction_prompt = (
        "Extract subject-predicate-object triples from the text below. "
        "Return ONLY a JSON array of objects with keys: subject, predicate, object. "
        "Example: [{\"subject\": \"Acme Corp\", \"predicate\": \"founded_by\", \"object\": \"Jane Smith\"}]\n\n"
        f"Text: {text[:1500]}\n\nJSON:"
    )

    triples = []
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(
                settings.RAY_LLM_ENDPOINT,
                json={
                    "model": settings.LLM_MODEL,
                    "messages": [{"role": "user", "content": extraction_prompt}],
                    "temperature": 0.0,
                    "max_tokens": 1024,
                    "stream": False,
                },
            )
            resp.raise_for_status()
            import json

            raw = resp.json()["choices"][0]["message"]["content"]
            # Try to find JSON array in the response
            start = raw.find("[")
            end = raw.rfind("]") + 1
            if start >= 0 and end > start:
                triples = json.loads(raw[start:end])
    except Exception as e:
        logger.warning(f"  ⚠️  Entity extraction failed: {e}")

    if not triples:
        logger.info("  ⏭️  No triples extracted — inserting document node only")
        triples = [{"subject": filename, "predicate": "contains", "object": text[:200]}]

    # Merge into Neo4j
    driver = GraphDatabase.driver(
        settings.NEO4J_URI,
        auth=(settings.NEO4J_USER, settings.NEO4J_PASSWORD),
    )
    with driver.session() as session:
        for t in triples:
            try:
                session.run(
                    """
                    MERGE (s:Entity {name: $subject})
                    SET s.tenant_id = $tenant_id
                    MERGE (o:Entity {name: $object})
                    SET o.tenant_id = $tenant_id
                    MERGE (s)-[r:RELATES {type: $predicate}]->(o)
                    SET r.source = $filename
                    """,
                    subject=str(t.get("subject", "")),
                    object=str(t.get("object", "")),
                    predicate=str(t.get("predicate", "")),
                    filename=filename,
                    tenant_id=tenant_id,
                )
            except Exception as e:
                logger.warning(f"  ⚠️  Neo4j merge failed for triple: {e}")
    driver.close()
    logger.info(f"  ✅ Neo4j: merged {len(triples)} triples")


# ── Multimodal document ingestion ───────────────────────────────────────────
async def ingest_image_file(filepath: str, tenant_id: str = "default"):
    """Ingest a standalone image file via Gemini multimodal embedding."""
    filename = os.path.basename(filepath)
    with open(filepath, "rb") as f:
        image_bytes = f.read()

    ext = os.path.splitext(filepath)[1].lower()
    mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp"}
    mime_type = mime_map.get(ext, "image/png")

    logger.info(f"\n🖼️  Ingesting image: {filename} ({len(image_bytes)} bytes, tenant={tenant_id})")

    # Embed image via Gemini
    embedding = await embed_image_gemini(image_bytes, mime_type)
    chunk = {"text": filename, "content_type": "image", "mime_type": mime_type}
    upsert_qdrant_multimodal([chunk], [embedding], filename, tenant_id=tenant_id)
    logger.info(f"  Image embedded ({len(embedding)} dims) and indexed")


# ── Main ─────────────────────────────────────────────────────────────────────
async def ingest_document(text: str, filename: str, tenant_id: str = "default"):
    if not text or not text.strip():
        logger.warning(f"  Skipping {filename} — empty content")
        return

    logger.info(f"\n📄 Ingesting: {filename} ({len(text)} chars, tenant={tenant_id})")

    # 1. Chunk
    chunks = chunk_text(text)
    logger.info(f"  Chunked into {len(chunks)} pieces")

    # 2. Embed — use Gemini when multimodal is enabled (shared vector space)
    if MULTIMODAL_ENABLED:
        logger.info(f"  Embedding {len(chunks)} chunks via Gemini...")
        embeddings = await embed_texts_gemini(chunks)
        chunk_dicts = [{"text": c, "content_type": "text"} for c in chunks]
        upsert_qdrant_multimodal(chunk_dicts, embeddings, filename, tenant_id=tenant_id)
    else:
        logger.info(f"  Embedding {len(chunks)} chunks via Ollama...")
        embeddings = await embed_texts(chunks)
        upsert_qdrant(chunks, embeddings, filename, tenant_id=tenant_id)

    # 3. Extract entities → Neo4j
    logger.info("  Extracting entities via LLM...")
    await create_neo4j_entities(text, filename, tenant_id=tenant_id)


def collect_files(path: str) -> list[str]:
    """Collect all supported files from a path (file or directory)."""
    if os.path.isfile(path):
        ext = os.path.splitext(path)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            logger.error(f"Unsupported file type: {ext}")
            logger.error(f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
            sys.exit(1)
        return [path]

    if os.path.isdir(path):
        files = []
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(glob.glob(os.path.join(path, f"**/*{ext}"), recursive=True))
        files.sort()
        return files

    logger.error(f"Path not found: {path}")
    sys.exit(1)


async def main():
    parser = argparse.ArgumentParser(
        description="Local document ingestion (supports .txt, .md, .pdf, .docx)",
        epilog="Examples:\n"
               "  python scripts/ingest_local.py report.pdf\n"
               "  python scripts/ingest_local.py ./docs/          # ingest whole directory\n"
               "  python scripts/ingest_local.py --sample         # built-in sample data\n",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("path", nargs="?", help="Path to a file (.txt/.md/.pdf/.docx) or directory")
    parser.add_argument("--sample", action="store_true", help="Ingest built-in sample data")
    parser.add_argument("--tenant-id", default="default", help="Tenant ID to tag data with (default: 'default')")
    args = parser.parse_args()
    tenant_id = args.tenant_id

    if args.sample:
        logger.info("=" * 55)
        logger.info(f"  Ingesting sample documents (tenant={tenant_id})...")
        logger.info("=" * 55)
        for doc in SAMPLE_DOCUMENTS:
            await ingest_document(doc["text"], doc["filename"], tenant_id=tenant_id)
    elif args.path:
        files = collect_files(args.path)
        if not files:
            logger.error(f"No supported files found in: {args.path}")
            logger.error(f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
            sys.exit(1)

        logger.info("=" * 55)
        logger.info(f"  Ingesting {len(files)} file(s) (tenant={tenant_id})...")
        logger.info("=" * 55)
        for filepath in files:
            try:
                ext = os.path.splitext(filepath)[1].lower()
                if ext in IMAGE_EXTENSIONS and MULTIMODAL_ENABLED:
                    await ingest_image_file(filepath, tenant_id=tenant_id)
                else:
                    text = parse_file(filepath)
                    await ingest_document(text, os.path.basename(filepath), tenant_id=tenant_id)
            except Exception as e:
                logger.error(f"  Failed to process {filepath}: {e}")
    else:
        parser.print_help()
        sys.exit(1)

    logger.info("\n" + "=" * 55)
    logger.info("  ✅  Ingestion complete!")
    logger.info("=" * 55)


if __name__ == "__main__":
    asyncio.run(main())
